import os
from pathlib import Path
from dotenv import load_dotenv
from openai import OpenAI
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json
from langgraph.graph import StateGraph, START, END
from typing import TypedDict

# Load environment variables
load_dotenv()

# Initialize NVIDIA LLM Client
client = OpenAI(
    base_url="https://integrate.api.nvidia.com/v1",
    api_key=os.getenv("NVIDIA_API_KEY")
)

# Define the state structure
class ResumeGenerationState(TypedDict):
    cv_content: str
    jd_content: str
    resume_data: dict
    docx_path: str
    pdf_path: str
    error: str | None

# Node 1: Load CV and JD files
def load_files(state: ResumeGenerationState) -> ResumeGenerationState:
    """Load CV.md and JD.txt files from the repo"""
    try:
        cv_path = Path("cv.md")
        jd_path = Path("jd.txt")
        
        if not cv_path.exists():
            return {**state, "error": f"CV file not found at {cv_path}"}
        if not jd_path.exists():
            return {**state, "error": f"JD file not found at {jd_path}"}
        
        with open(cv_path, "r", encoding="utf-8") as f:
            cv_content = f.read()
        with open(jd_path, "r", encoding="utf-8") as f:
            jd_content = f.read()
        
        print(f"✓ Loaded CV ({len(cv_content)} chars) and JD ({len(jd_content)} chars)")
        
        return {
            **state,
            "cv_content": cv_content,
            "jd_content": jd_content,
            "error": None
        }
    except Exception as e:
        return {**state, "error": f"Error loading files: {str(e)}"}

# Node 2: Generate complete resume data using LLM
def generate_resume_data(state: ResumeGenerationState) -> ResumeGenerationState:
    """Use LLM to generate structured resume data tailored to job description"""
    if state.get("error"):
        return state
    
    try:
        prompt = f"""Extract and tailor complete resume data from the CV to match the job description. Return ONLY a valid JSON object with NO markdown, NO code blocks, NO explanations.

{{
    "professional_summary": "2-3 sentences tailored to JD keywords highlighting relevant skills and experience",
    "experience": [
        {{
            "company": "Company Name",
            "title": "Job Title",
            "dates": "Month Year – Month Year",
            "achievements": [
                "Achievement 1 with quantified results",
                "Achievement 2 with quantified results",
                "Achievement 3 with quantified results",
                "Achievement 4 with quantified results",
                "Achievement 5 if available"
            ]
        }},
        {{
            "company": "Company Name",
            "title": "Job Title",
            "dates": "Month Year – Month Year",
            "achievements": [
                "Achievement 1 with quantified results",
                "Achievement 2 with quantified results",
                "Achievement 3 with quantified results",
                "Achievement 4 with quantified results"
            ]
        }}
    ],
    "projects": [
        {{
            "name": "Project Name",
            "description": "Brief description with relevant technologies",
            "achievement": "Key achievement with metrics"
        }},
        {{
            "name": "Project Name",
            "description": "Brief description with relevant technologies",
            "achievement": "Key achievement with metrics"
        }}
    ],
    "skills": [
        {{
            "category": "Programming Languages",
            "items": "List of programming languages relevant to JD"
        }},
        {{
            "category": "AI & LLM",
            "items": "List of AI/LLM skills"
        }},
        {{
            "category": "Cloud & Infrastructure",
            "items": "List of cloud/infrastructure skills"
        }},
        {{
            "category": "DevOps & Tools",
            "items": "List of DevOps tools"
        }},
        {{
            "category": "Databases",
            "items": "List of databases"
        }}
    ],
    "achievements": [
        "Achievement 1",
        "Achievement 2",
        "Achievement 3",
        "Achievement 4",
        "Achievement 5"
    ]
}}

CV CONTENT:
{state['cv_content']}

JOB DESCRIPTION:
{state['jd_content']}

IMPORTANT RULES:
1. Tailor ALL content to match job description keywords
2. Use quantified metrics (e.g., "32x faster", "25% reduction", "1M+ users")
3. Prioritize JD requirements in skills and achievements
4. Keep person name: Prateek Khandelwal
5. Keep contact: +91 6375954380 | contact.prateekcse1@gmail.com | GitHub | LinkedIn | LeetCode
6. Return ONLY the JSON object - no markdown backticks, no explanation
7. Valid JSON format only

Return the JSON object:"""

        print("\n🔄 Generating tailored resume data with NVIDIA Nemotron LLM...")
        
        response = client.chat.completions.create(
            model="nvidia/llama-3.3-nemotron-super-49b-v1",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.6,
            top_p=0.95,
            max_tokens=4096,
            frequency_penalty=0,
            presence_penalty=0,
            stream=False
        )
        
        response_text = response.choices[0].message.content
        resume_data = json.loads(response_text)
        print(f"✓ Generated tailored resume data successfully")
        
        return {
            **state,
            "resume_data": resume_data,
            "error": None
        }
    except json.JSONDecodeError as e:
        return {**state, "error": f"Error parsing LLM JSON response: {str(e)}"}
    except Exception as e:
        return {**state, "error": f"Error generating resume data: {str(e)}"}

# Node 3: Create DOCX with your template structure
def create_resume_docx(state: ResumeGenerationState) -> ResumeGenerationState:
    """Create DOCX following your resume template structure and theme"""
    if state.get("error"):
        return state
    
    try:
        print("\n🔄 Creating resume document with your template theme...")
        
        doc = Document()
        
        # Set margins like your template
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
        
        data = state["resume_data"]
        
        # 1. NAME - Bold, Centered, Large
        name_para = doc.add_paragraph("Prateek Khandelwal")
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name_para.runs[0]
        name_run.font.size = Pt(18)
        name_run.font.bold = True
        name_run.font.color.rgb = RGBColor(31, 31, 31)
        
        # 2. CONTACT INFO - Centered, Blue links
        contact_para = doc.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        contact_items = [
            ("+91 6375954380", "tel:+916375954380"),
            ("contact.prateekcse1@gmail.com", "mailto:contact.prateekcse1@gmail.com"),
            ("GitHub", "https://github.com/prateek-khandelwal"),
            ("LinkedIn", "https://linkedin.com/in/prateek-khandelwal"),
            ("LeetCode", "https://leetcode.com/prateek-khandelwal"),
        ]
        
        for i, (text, url) in enumerate(contact_items):
            if i > 0:
                contact_para.add_run(" | ").font.color.rgb = RGBColor(0, 102, 204)
            run = contact_para.add_run(text)
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0, 102, 204)
        
        # Add horizontal line after header
        line_para = doc.add_paragraph("_" * 95)
        line_run = line_para.runs[0]
        line_run.font.color.rgb = RGBColor(0, 0, 0)
        
        # 3. PROFESSIONAL SUMMARY with proper line
        summary_heading = doc.add_paragraph("PROFESSIONAL SUMMARY")
        summary_heading_run = summary_heading.runs[0]
        summary_heading_run.font.bold = True
        summary_heading_run.font.size = Pt(11)
        summary_heading_run.font.color.rgb = RGBColor(31, 31, 31)
        summary_heading_run.underline = False
        summary_heading.paragraph_format.space_after = Pt(2)
        
        # Add bottom border line
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        pPr = summary_heading._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '12')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '000000')
        pBdr.append(bottom)
        pPr.append(pBdr)
        
        summary_para = doc.add_paragraph(data.get("professional_summary", ""))
        summary_para.paragraph_format.space_before = Pt(3)
        for run in summary_para.runs:
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(31, 31, 31)
        
        # 4. EXPERIENCE with proper line
        exp_heading = doc.add_paragraph("EXPERIENCE")
        exp_heading_run = exp_heading.runs[0]
        exp_heading_run.font.bold = True
        exp_heading_run.font.size = Pt(11)
        exp_heading_run.font.color.rgb = RGBColor(31, 31, 31)
        exp_heading.paragraph_format.space_before = Pt(10)
        exp_heading.paragraph_format.space_after = Pt(2)
        
        # Add bottom border line
        pPr = exp_heading._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '12')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '000000')
        pBdr.append(bottom)
        pPr.append(pBdr)
        
        for job in data.get("experience", []):
            # Company and dates
            company_para = doc.add_paragraph()
            company_run = company_para.add_run(job["company"])
            company_run.font.bold = True
            company_run.font.size = Pt(11)
            company_run.font.color.rgb = RGBColor(31, 31, 31)
            
            # Add dates on the right
            tabs = company_para.paragraph_format.tab_stops
            tabs.add_tab_stop(Inches(5.5))
            company_para.text = job["company"] + "\t" + job["dates"]
            
            # Title
            title_para = doc.add_paragraph(job["title"])
            title_run = title_para.runs[0]
            title_run.font.italic = True
            title_run.font.size = Pt(11)
            title_run.font.color.rgb = RGBColor(31, 31, 31)
            
            # Achievements as bullets
            for achievement in job.get("achievements", []):
                bullet_para = doc.add_paragraph(achievement, style='List Bullet')
                bullet_para.paragraph_format.left_indent = Inches(0.25)
                for run in bullet_para.runs:
                    run.font.size = Pt(11)
                    run.font.color.rgb = RGBColor(31, 31, 31)
        
        # 5. PROJECTS with proper line
        proj_heading = doc.add_paragraph("PROJECTS")
        proj_heading_run = proj_heading.runs[0]
        proj_heading_run.font.bold = True
        proj_heading_run.font.size = Pt(11)
        proj_heading_run.font.color.rgb = RGBColor(31, 31, 31)
        proj_heading.paragraph_format.space_before = Pt(10)
        proj_heading.paragraph_format.space_after = Pt(2)
        
        # Add bottom border line
        pPr = proj_heading._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '12')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '000000')
        pBdr.append(bottom)
        pPr.append(pBdr)
        
        for project in data.get("projects", []):
            # Project name
            proj_name = doc.add_paragraph(project["name"])
            proj_name_run = proj_name.runs[0]
            proj_name_run.font.bold = True
            proj_name_run.font.size = Pt(11)
            proj_name_run.font.color.rgb = RGBColor(31, 31, 31)
            
            # Description
            desc_para = doc.add_paragraph(project["description"], style='List Bullet')
            desc_para.paragraph_format.left_indent = Inches(0.25)
            for run in desc_para.runs:
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(31, 31, 31)
            
            # Achievement
            ach_para = doc.add_paragraph(project["achievement"], style='List Bullet')
            ach_para.paragraph_format.left_indent = Inches(0.25)
            for run in ach_para.runs:
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(31, 31, 31)
        
        # 6. SKILLS with proper line
        skills_heading = doc.add_paragraph("SKILLS")
        skills_heading_run = skills_heading.runs[0]
        skills_heading_run.font.bold = True
        skills_heading_run.font.size = Pt(11)
        skills_heading_run.font.color.rgb = RGBColor(31, 31, 31)
        skills_heading.paragraph_format.space_before = Pt(10)
        skills_heading.paragraph_format.space_after = Pt(2)
        
        # Add bottom border line
        pPr = skills_heading._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '12')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '000000')
        pBdr.append(bottom)
        pPr.append(pBdr)
        
        for skill in data.get("skills", []):
            skill_para = doc.add_paragraph()
            
            category_run = skill_para.add_run(f"{skill['category']}: ")
            category_run.font.bold = True
            category_run.font.size = Pt(11)
            category_run.font.color.rgb = RGBColor(31, 31, 31)
            
            items_run = skill_para.add_run(skill['items'])
            items_run.font.size = Pt(11)
            items_run.font.color.rgb = RGBColor(31, 31, 31)
        
        # 7. ACHIEVEMENTS & CERTIFICATIONS with proper line
        ach_heading = doc.add_paragraph("ACHIEVEMENTS & CERTIFICATIONS")
        ach_heading_run = ach_heading.runs[0]
        ach_heading_run.font.bold = True
        ach_heading_run.font.size = Pt(11)
        ach_heading_run.font.color.rgb = RGBColor(31, 31, 31)
        ach_heading.paragraph_format.space_before = Pt(10)
        ach_heading.paragraph_format.space_after = Pt(2)
        
        # Add bottom border line
        pPr = ach_heading._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '12')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '000000')
        pBdr.append(bottom)
        pPr.append(pBdr)
        
        for achievement in data.get("achievements", []):
            ach_para = doc.add_paragraph(achievement, style='List Bullet')
            ach_para.paragraph_format.left_indent = Inches(0.25)
            for run in ach_para.runs:
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(31, 31, 31)
        
        # Save document
        output_dir = Path("generated_resumes")
        output_dir.mkdir(exist_ok=True)
        
        docx_path = output_dir / "tailored_resume.docx"
        doc.save(str(docx_path))
        
        print(f"✓ Created resume DOCX at {docx_path}")
        
        return {
            **state,
            "docx_path": str(docx_path),
            "error": None
        }
    except Exception as e:
        return {**state, "error": f"Error creating DOCX: {str(e)}"}

# Node 4: Convert DOCX to PDF
def convert_to_pdf(state: ResumeGenerationState) -> ResumeGenerationState:
    """Convert DOCX to PDF using LibreOffice"""
    if state.get("error"):
        return state
    
    try:
        print("\n🔄 Converting to PDF...")
        
        import subprocess
        output_dir = Path("generated_resumes")
        docx_path = state["docx_path"]
        
        soffice_paths = [
            "C:\\Program Files\\LibreOffice\\program\\soffice.exe",
            "C:\\Program Files (x86)\\LibreOffice\\program\\soffice.exe",
        ]
        
        soffice_path = None
        for path in soffice_paths:
            if os.path.exists(path):
                soffice_path = path
                break
        
        if soffice_path:
            cmd = [
                soffice_path,
                "--headless",
                "--convert-to", "pdf",
                "--outdir", str(output_dir),
                docx_path
            ]
            subprocess.run(cmd, capture_output=True, timeout=60)
            
            pdf_path = output_dir / "tailored_resume.pdf"
            if pdf_path.exists():
                print(f"✓ PDF saved to {pdf_path}")
                return {**state, "pdf_path": str(pdf_path), "error": None}
        
        print("⚠️  LibreOffice not found. DOCX created successfully.")
        print("   To convert: Open DOCX → Export as PDF")
        return {**state, "pdf_path": docx_path, "error": None}
        
    except Exception as e:
        print(f"⚠️  PDF conversion warning: {str(e)}")
        return {**state, "pdf_path": state["docx_path"], "error": None}

# Build the graph
def build_graph():
    """Build the LangGraph workflow"""
    graph = StateGraph(ResumeGenerationState)
    
    graph.add_node("load_files", load_files)
    graph.add_node("generate_data", generate_resume_data)
    graph.add_node("create_docx", create_resume_docx)
    graph.add_node("convert_pdf", convert_to_pdf)
    
    graph.add_edge(START, "load_files")
    graph.add_edge("load_files", "generate_data")
    graph.add_edge("generate_data", "create_docx")
    graph.add_edge("create_docx", "convert_pdf")
    graph.add_edge("convert_pdf", END)
    
    return graph.compile()

# Main execution
def main():
    """Run the resume generation workflow"""
    print("=" * 70)
    print("🚀 CVbyAI - ATS-Optimized Resume Generator")
    print("   Tailored content + Your template structure")
    print("=" * 70)
    
    initial_state: ResumeGenerationState = {
        "cv_content": "",
        "jd_content": "",
        "resume_data": {},
        "docx_path": "",
        "pdf_path": "",
        "error": None
    }
    
    graph = build_graph()
    final_state = graph.invoke(initial_state)
    
    print("\n" + "=" * 70)
    if final_state.get("error"):
        print(f"❌ Error: {final_state['error']}")
    else:
        print("✅ Resume generated successfully!")
        print(f"\n📄 DOCX: {final_state['docx_path']}")
        if final_state.get('pdf_path') and final_state['pdf_path'] != final_state['docx_path']:
            print(f"📕 PDF: {final_state['pdf_path']}")
        print("\n✨ LLM-tailored + Your template structure!")
    print("=" * 70)

if __name__ == "__main__":
    main()
